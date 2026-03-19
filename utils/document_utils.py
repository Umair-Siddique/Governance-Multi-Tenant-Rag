"""
Document text extraction and preprocessing utilities.
Handles PDF, DOCX, CSV, and related formats with normalization and cleaning.
When documents contain images, Tesseract OCR is used to extract text from those images (if available).
CSV uses column-aware parsing and row-group chunking for Supabase registry + Pinecone search.
"""
import csv
import io
import re
import zipfile
import unicodedata
from pathlib import Path
from typing import BinaryIO, Iterator, List, Optional, Tuple, Union

from pypdf import PdfReader
from docx import Document

try:
    import pytesseract
    from PIL import Image
    _OCR_AVAILABLE = True
except ImportError:
    _OCR_AVAILABLE = False
    pytesseract = None
    Image = None


# --- Supported file extensions ---
SUPPORTED_PDF_EXTENSIONS = {".pdf"}
SUPPORTED_DOCX_EXTENSIONS = {".docx"}
SUPPORTED_CSV_EXTENSIONS = {".csv"}
# Image formats processed via Tesseract OCR (PIL-openable)
SUPPORTED_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".tif", ".webp"}


# --- Unicode categories to strip (control, format, surrogate, private-use) ---
_CATEGORIES_TO_STRIP = ("Cc", "Cf", "Cs", "Co", "Cn")


def _ocr_image(image_input: Union["Image.Image", bytes], tesseract_cmd: Optional[str] = None) -> str:
    """
    Run Tesseract OCR on a single image. Returns extracted text or empty string if OCR is unavailable or fails.
    """
    if not _OCR_AVAILABLE or not pytesseract or not Image:
        return ""
    try:
        if tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
        if isinstance(image_input, bytes):
            img = Image.open(io.BytesIO(image_input))
        else:
            img = image_input
        if img.mode not in ("L", "RGB", "RGBA"):
            img = img.convert("RGB")
        text = pytesseract.image_to_string(img)
        return (text or "").strip()
    except Exception:
        return ""


def extract_text_from_image(
    source: Union[str, Path, bytes, BinaryIO],
    tesseract_cmd: Optional[str] = None,
) -> str:
    """
    Extract text from an image file using Tesseract OCR.
    Same cleaning/normalization as other formats is applied later via preprocess_text.

    Args:
        source: File path, bytes, or file-like object.
        tesseract_cmd: Optional path to tesseract executable.

    Returns:
        Extracted text as a single string.
    """
    if not _OCR_AVAILABLE or not Image:
        return ""
    try:
        if isinstance(source, (str, Path)):
            img = Image.open(str(source))
        elif isinstance(source, bytes):
            img = Image.open(io.BytesIO(source))
        elif hasattr(source, "read"):
            source.seek(0)
            img = Image.open(io.BytesIO(source.read()))
        else:
            raise TypeError("source must be path, bytes, or file-like object")
        return _ocr_image(img, tesseract_cmd)
    except Exception:
        return ""


def _iter_pdf_images(reader: PdfReader) -> Iterator[Union["Image.Image", bytes]]:
    """Yield PIL Images or raw bytes for each embedded image in the PDF (for OCR). Skips images that fail to decode (e.g. invalid lookup table, unsupported color space)."""
    if not _OCR_AVAILABLE or not Image:
        return
    for page in reader.pages:
        images_container = getattr(page, "images", None)
        if images_container is None:
            continue
        # pypdf: iterate by key so one bad image doesn't break the whole page (see pypdf docs on error handling)
        keys = getattr(images_container, "keys", None)
        if keys is not None:
            names = list(keys())
        else:
            try:
                names = list(range(len(images_container)))
            except Exception:
                continue
        for name in names:
            try:
                img_obj = images_container[name]
            except Exception:
                continue
            try:
                pil_img = getattr(img_obj, "image", None)
                if pil_img is not None:
                    yield pil_img
                    continue
                data = getattr(img_obj, "data", None)
                if data:
                    yield Image.open(io.BytesIO(data))
            except Exception:
                # Skip images that fail to decode (e.g. Invalid Lookup Table, unsupported color space)
                continue


def _iter_docx_images(source: Union[str, Path, bytes, BinaryIO]) -> Iterator[bytes]:
    """Yield raw image bytes for each embedded image in the DOCX (e.g. word/media/*)."""
    if isinstance(source, (str, Path)):
        fp = open(source, "rb")
        try:
            z = zipfile.ZipFile(fp, "r")
        except zipfile.BadZipFile:
            fp.close()
            return
    elif isinstance(source, bytes):
        z = zipfile.ZipFile(io.BytesIO(source), "r")
        fp = None
    elif hasattr(source, "read"):
        try:
            source.seek(0)
            z = zipfile.ZipFile(io.BytesIO(source.read()), "r")
        except Exception:
            return
        fp = None
    else:
        return
    try:
        for name in z.namelist():
            if name.startswith("word/media/") and name.rstrip("/") != "word/media":
                try:
                    yield z.read(name)
                except Exception:
                    continue
    finally:
        z.close()
        if fp is not None:
            fp.close()


def extract_text_from_pdf(
    source: Union[str, Path, bytes, BinaryIO],
    password: Optional[str] = None,
    tesseract_cmd: Optional[str] = None,
) -> str:
    """
    Extract raw text from a PDF file. When the PDF contains embedded images,
    Tesseract OCR is used to extract text from those images (if available).

    Args:
        source: File path, bytes, or file-like object.
        password: Optional password for protected PDFs.
        tesseract_cmd: Optional path to tesseract executable (e.g. on Windows). If None, system PATH is used.

    Returns:
        Extracted text as a single string.
    """
    if isinstance(source, (str, Path)):
        reader = PdfReader(str(source))
    elif isinstance(source, bytes):
        reader = PdfReader(io.BytesIO(source))
    elif hasattr(source, "read"):
        reader = PdfReader(source)
    else:
        raise TypeError("source must be path, bytes, or file-like object")

    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)

    # OCR text from embedded images when Tesseract is available (skip images that fail to decode)
    if _OCR_AVAILABLE:
        for img in _iter_pdf_images(reader):
            try:
                ocr_text = _ocr_image(img, tesseract_cmd)
                if ocr_text:
                    parts.append(ocr_text)
            except Exception:
                continue

    raw = "\n".join(parts)
    return raw if raw else ""


def extract_text_from_docx(
    source: Union[str, Path, bytes, BinaryIO],
    tesseract_cmd: Optional[str] = None,
) -> str:
    """
    Extract raw text from a DOCX file. When the document contains embedded images,
    Tesseract OCR is used to extract text from those images (if available).

    Args:
        source: File path, bytes, or file-like object.
        tesseract_cmd: Optional path to tesseract executable. If None, system PATH is used.

    Returns:
        Extracted text as a single string.
    """
    if isinstance(source, (str, Path)):
        doc = Document(str(source))
        source_for_images = source
    elif isinstance(source, bytes):
        doc = Document(io.BytesIO(source))
        source_for_images = source
    elif hasattr(source, "read"):
        data = source.read()
        doc = Document(io.BytesIO(data))
        source_for_images = data
    else:
        raise TypeError("source must be path, bytes, or file-like object")

    parts = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)

    # OCR text from embedded images when Tesseract is available
    if _OCR_AVAILABLE:
        for img_bytes in _iter_docx_images(source_for_images):
            ocr_text = _ocr_image(img_bytes, tesseract_cmd)
            if ocr_text:
                parts.append(ocr_text)

    raw = "\n".join(parts)
    return raw if raw else ""


def remove_control_and_special_chars(text: str) -> str:
    """
    Remove control characters, formatting chars, and other problematic Unicode.
    """
    if not text:
        return ""
    return "".join(c for c in text if unicodedata.category(c) not in _CATEGORIES_TO_STRIP)


def remove_extra_whitespace(text: str) -> str:
    """Collapse multiple spaces/newlines and strip edges."""
    if not text:
        return ""
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def remove_page_numbers(text: str) -> str:
    """
    Remove common page number patterns (e.g. "Page 1 of 10", "- 5 -", "5" on its own line).
    """
    if not text:
        return ""
    patterns = [
        r"\b[Pp]age\s+\d+\s+[Oo]f\s+\d+\b",
        r"\b\d+\s+[Oo]f\s+\d+\b",
        r"^\s*[-–—]\s*\d+\s*[-–—]\s*$",
        r"^\s*\d+\s*$",
    ]
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        for pat in patterns:
            line = re.sub(pat, "", line, flags=re.IGNORECASE)
        line = line.strip()
        if line:
            cleaned.append(line)
    return "\n".join(cleaned)


def remove_repeated_sentences(text: str) -> str:
    """
    Remove consecutive duplicate lines/sentences (common OCR artifact).
    """
    if not text:
        return ""
    lines = text.split("\n")
    result = []
    prev = None
    for line in lines:
        line_stripped = line.strip()
        if line_stripped and line_stripped != prev:
            result.append(line)
            prev = line_stripped
    return "\n".join(result)


def remove_ocr_noise(text: str) -> str:
    """
    Remove typical OCR artifacts: stray punctuation, repeated chars, etc.
    """
    if not text:
        return ""
    # Replace repeated non-word chars (e.g. "...", "---")
    text = re.sub(r"([^\w\s])\1{2,}", r"\1", text)
    # Remove single stray punctuation on its own
    text = re.sub(r"\n\s*[^\w\s]\s*\n", "\n", text)
    return text


def remove_headers_footers(text: str) -> str:
    """
    Heuristic: remove lines that look like headers/footers (short, repeated, or
    containing typical header/footer keywords).
    """
    if not text:
        return ""
    lines = text.split("\n")
    header_footer_keywords = {
        "confidential", "draft", "copyright", "©", "page", "©",
        "all rights reserved", "proprietary", "internal use only",
    }
    result = []
    for line in lines:
        s = line.strip().lower()
        if len(s) < 3:
            continue
        if any(kw in s for kw in header_footer_keywords) and len(s) < 80:
            continue
        result.append(line)
    return "\n".join(result)


def normalize_unicode(text: str) -> str:
    """
    Normalize Unicode (e.g. NFD → NFC) for consistent handling.
    """
    if not text:
        return ""
    return unicodedata.normalize("NFKC", text)


def preprocess_text(
    text: str,
    remove_control_chars: bool = True,
    normalize_unicode_text: bool = True,
    collapse_whitespace: bool = True,
    remove_page_nums: bool = True,
    remove_duplicate_lines: bool = True,
    remove_ocr_artifacts: bool = True,
    remove_header_footer: bool = False,
) -> str:
    """
    Preprocess extracted text: clean and normalize.

    Args:
        text: Raw extracted text.
        remove_control_chars: Strip control/format/special Unicode.
        normalize_unicode_text: NFKC normalization.
        collapse_whitespace: Collapse multiple spaces/newlines.
        remove_page_nums: Remove page number patterns.
        remove_duplicate_lines: Remove consecutive duplicate lines.
        remove_ocr_artifacts: Remove OCR noise patterns.
        remove_header_footer: Heuristic removal of header/footer lines.

    Returns:
        Cleaned and normalized text.
    """
    if not text:
        return ""

    if remove_control_chars:
        text = remove_control_and_special_chars(text)
    if normalize_unicode_text:
        text = normalize_unicode(text)
    if remove_page_nums:
        text = remove_page_numbers(text)
    if remove_duplicate_lines:
        text = remove_repeated_sentences(text)
    if remove_ocr_artifacts:
        text = remove_ocr_noise(text)
    if remove_header_footer:
        text = remove_headers_footers(text)
    if collapse_whitespace:
        text = remove_extra_whitespace(text)

    return text


def extract_and_preprocess(
    source: Union[str, Path, bytes, BinaryIO],
    file_ext: Optional[str] = None,
    password: Optional[str] = None,
    preprocess: bool = True,
    tesseract_cmd: Optional[str] = None,
    **preprocess_kwargs,
) -> str:
    """
    Extract text from a document and optionally preprocess it.
    When the document contains images, Tesseract OCR is used to extract text from them (if available).

    Args:
        source: File path, bytes, or file-like object.
        file_ext: Extension (e.g. '.pdf', '.docx'). Inferred from path if possible.
        password: Optional PDF password.
        preprocess: Whether to run preprocessing.
        tesseract_cmd: Optional path to tesseract executable (e.g. Windows: r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'). On production (e.g. Render) leave unset to use system tesseract from PATH.
        **preprocess_kwargs: Passed to preprocess_text().

    Returns:
        Extracted (and optionally preprocessed) text.
    """
    if file_ext is None and isinstance(source, (str, Path)):
        file_ext = Path(source).suffix.lower()
    if file_ext is None and hasattr(source, "name"):
        file_ext = Path(getattr(source, "name", "") or "").suffix.lower()
    if not file_ext:
        file_ext = ".pdf"  # fallback

    file_ext = file_ext.lower() if file_ext.startswith(".") else f".{file_ext}".lower()

    if file_ext in SUPPORTED_PDF_EXTENSIONS:
        raw = extract_text_from_pdf(source, password=password, tesseract_cmd=tesseract_cmd)
    elif file_ext in SUPPORTED_DOCX_EXTENSIONS:
        raw = extract_text_from_docx(source, tesseract_cmd=tesseract_cmd)
    elif file_ext in SUPPORTED_IMAGE_EXTENSIONS:
        raw = extract_text_from_image(source, tesseract_cmd=tesseract_cmd)
    else:
        raise ValueError(
            f"Unsupported file type: {file_ext}. Supported: PDF, DOCX, CSV, and images (e.g. JPG, PNG)"
        )

    if preprocess:
        return preprocess_text(raw, **preprocess_kwargs)
    return raw


def get_supported_extensions() -> set:
    """Return all supported file extensions for extraction (PDF, DOCX, CSV, images)."""
    return (
        SUPPORTED_PDF_EXTENSIONS
        | SUPPORTED_DOCX_EXTENSIONS
        | SUPPORTED_CSV_EXTENSIONS
        | SUPPORTED_IMAGE_EXTENSIONS
    )


def is_supported(filename_or_ext: str) -> bool:
    """Check if the given filename or extension is supported."""
    ext = filename_or_ext.lower()
    if not ext.startswith("."):
        ext = Path(ext).suffix.lower() if "." in ext else f".{ext}"
    return ext in get_supported_extensions()


# --- CSV: column-aware parsing and row-group chunking ---
def _normalize_csv_cell(value: str) -> str:
    """Normalize a single CSV cell: strip and collapse internal whitespace."""
    if not value:
        return ""
    return " ".join((value or "").strip().split())


def parse_csv(
    source: Union[str, Path, bytes, BinaryIO],
    max_rows: Optional[int] = None,
    encoding: str = "utf-8",
    delimiter: Optional[str] = None,
) -> Tuple[List[str], List[dict]]:
    """
    Parse CSV with column-aware handling. Returns column names and list of row dicts.

    Args:
        source: File path, bytes, or file-like object.
        max_rows: If set, only read this many data rows (after header).
        encoding: Text encoding (default utf-8).
        delimiter: CSV delimiter; if None, sniffed (comma or tab).

    Returns:
        (columns, rows) where columns is list of header strings, rows is list of dicts
        with normalized string values.
    """
    if isinstance(source, (str, Path)):
        with open(source, "r", encoding=encoding, newline="") as f:
            text = f.read()
    elif isinstance(source, bytes):
        text = source.decode(encoding)
    elif hasattr(source, "read"):
        raw = source.read()
        text = raw.decode(encoding) if isinstance(raw, bytes) else raw
    else:
        raise TypeError("source must be path, bytes, or file-like object")

    if delimiter is None:
        try:
            dialect = csv.Sniffer().sniff(text[:8192], delimiters=",\t;")
            delimiter = dialect.delimiter
        except (csv.Error, TypeError):
            delimiter = ","

    f = io.StringIO(text)
    reader = csv.DictReader(f, delimiter=delimiter)
    columns = [c for c in (reader.fieldnames or []) if c]
    if not columns:
        return [], []
    rows = []
    for i, row in enumerate(reader):
        if max_rows is not None and i >= max_rows:
            break
        normalized = {k: _normalize_csv_cell(str(row.get(k, ""))) for k in columns}
        rows.append(normalized)
    return columns, rows


def csv_rows_to_chunks(
    columns: List[str],
    rows: List[dict],
    chunk_size: int = 2000,
    rows_per_chunk: Optional[int] = None,
) -> List[str]:
    """
    Turn CSV rows into searchable text chunks for embedding.
    Each chunk is a text block of "column: value" lines for a group of rows.

    Args:
        columns: Column names.
        rows: List of row dicts (same keys as columns).
        chunk_size: Target max characters per chunk if rows_per_chunk not set.
        rows_per_chunk: If set, each chunk contains exactly this many rows (except last).

    Returns:
        List of chunk strings.
    """
    if not columns or not rows:
        return []

    def row_to_line(row: dict) -> str:
        parts = [f"{col}: {row.get(col, '')}" for col in columns]
        return " | ".join(parts)

    if rows_per_chunk is not None and rows_per_chunk > 0:
        out = []
        for i in range(0, len(rows), rows_per_chunk):
            group = rows[i : i + rows_per_chunk]
            out.append("\n".join(row_to_line(r) for r in group))
        return out

    chunks = []
    current_lines = []
    current_len = 0
    sep_len = 1

    for row in rows:
        line = row_to_line(row)
        line_len = len(line) + sep_len
        if current_lines and current_len + line_len > chunk_size:
            chunks.append("\n".join(current_lines))
            current_lines = [line]
            current_len = line_len
        else:
            current_lines.append(line)
            current_len += line_len

    if current_lines:
        chunks.append("\n".join(current_lines))
    return chunks


def generate_csv_summary_with_llm(
    columns: List[str],
    rows: List[dict],
    openai_client,
    max_sample_rows: int = 20,
    max_summary_chars: int = 2000,
) -> str:
    """
    Generate a short summary of the CSV using LLM (for csv_registry.summary).
    Uses only column names and a small sample of rows; no interpretation or advice.
    """
    if not openai_client or not columns:
        return ""
    sample = rows[:max_sample_rows]
    sample_text = "\n".join(
        " | ".join(str(r.get(c, "")) for c in columns) for r in sample
    )
    prompt = (
        "You are a data cataloguer. Based only on the column names and the sample rows below, "
        "write a brief factual summary (2–4 sentences) of what this dataset contains. "
        "Do not add interpretation, advice, or information not present in the data.\n\n"
        f"Columns: {', '.join(columns)}\n\nSample rows:\n{sample_text}"
    )
    try:
        resp = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=256,
        )
        summary = (resp.choices[0].message.content or "").strip()
        return summary[:max_summary_chars] if summary else ""
    except Exception:
        return ""


# --- Recursive text chunking ---
DEFAULT_CHUNK_SIZE = 2000
DEFAULT_CHUNK_OVERLAP = 200
DEFAULT_SEPARATORS = ["\n\n", "\n", ". ", ", ", " ", ""]


def recursive_chunk(
    text: str,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
    separators: Optional[list] = None,
) -> list[str]:
    """
    Split text into chunks using recursive splitting with overlap.

    Tries separators in order (paragraph → line → sentence → word → char).
    If a piece exceeds chunk_size, recursively splits it with the next separator.
    Overlap is applied between consecutive chunks.

    Args:
        text: Text to split.
        chunk_size: Target max characters per chunk.
        chunk_overlap: Character overlap between consecutive chunks.
        separators: List of separator strings (tried in order).

    Returns:
        List of chunk strings.
    """
    if not text or not text.strip():
        return []

    if separators is None:
        separators = list(DEFAULT_SEPARATORS)
    overlap = min(max(0, chunk_overlap), chunk_size - 1)

    def _split_recursive(t: str, sep_idx: int) -> list[str]:
        sep = separators[sep_idx] if sep_idx < len(separators) else ""
        if sep:
            raw = t.split(sep)
            parts = [p.strip() for p in raw if p.strip()]
        else:
            parts = list(t)

        chunks = []
        current = ""
        for part in parts:
            if not part:
                continue
            sep_str = sep if current and sep else ""
            candidate = (current + sep_str + part) if current else part

            if len(candidate) <= chunk_size:
                current = candidate
            else:
                if current:
                    chunks.append(current)
                if len(part) > chunk_size:
                    sub = _split_recursive(part, sep_idx + 1) if sep_idx + 1 < len(separators) else _split_by_char(part)
                    for i, c in enumerate(sub):
                        if i == 0 and overlap and chunks:
                            c = chunks[-1][-overlap:] + c
                        chunks.append(c)
                    current = ""
                else:
                    overlap_str = current[-overlap:] if len(current) >= overlap else current
                    current = overlap_str + sep_str + part

        if current:
            chunks.append(current)
        return chunks

    def _split_by_char(s: str) -> list[str]:
        out = []
        start = 0
        while start < len(s):
            end = min(start + chunk_size, len(s))
            out.append(s[start:end])
            start = end - overlap if end < len(s) else len(s)
        return out

    return _split_recursive(text.strip(), 0)
